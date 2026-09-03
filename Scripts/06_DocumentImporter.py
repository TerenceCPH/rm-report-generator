#!/usr/bin/env python3
"""Extract DOORS-ready Markdown and DOCX from paired DOCX+PDF documents."""

from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import subprocess
import sys
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.parts.image import ImagePart
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from env import DOCUMENT_IMPORTER_DIR, output_subdir

OUTPUT_FONT_NAME = "Arial"
OUTPUT_FONT_SIZE = Pt(12)
# Word border size is in eighths of a point (4 = 0.5pt)
OUTPUT_TABLE_BORDER_SZ = "4"
OUTPUT_TABLE_BORDER_COLOR = "000000"

# VML namespace (not in python-docx default nsmap)
VML_IMAGEDATA = "{urn:schemas-microsoft-com:vml}imagedata"
A_BLIP = qn("a:blip")
R_EMBED = qn("r:embed")
R_ID = qn("r:id")
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DGM_REL_IDS = "{http://schemas.openxmlformats.org/drawingml/2006/diagram}relIds"
W_DRAWING = qn("w:drawing")
W_R = qn("w:r")
# python-docx has no RT constant for the MS diagram drawing part
RT_DIAGRAM_DRAWING = (
    "http://schemas.microsoft.com/office/2007/relationships/diagramDrawing"
)
DGM_ATTR_TO_RT = {
    "dm": RT.DIAGRAM_DATA,
    "lo": RT.DIAGRAM_LAYOUT,
    "qs": RT.DIAGRAM_QUICK_STYLE,
    "cs": RT.DIAGRAM_COLORS,
}
DGM_RT_TO_PARTNAME = {
    RT.DIAGRAM_DATA: "/word/diagrams/data%d.xml",
    RT.DIAGRAM_LAYOUT: "/word/diagrams/layout%d.xml",
    RT.DIAGRAM_QUICK_STYLE: "/word/diagrams/quickStyle%d.xml",
    RT.DIAGRAM_COLORS: "/word/diagrams/colors%d.xml",
    RT_DIAGRAM_DRAWING: "/word/diagrams/drawing%d.xml",
}
SMARTART_PLACEHOLDER = "[SmartArt diagram]"

DEFAULT_INPUT_DIRS = (Path(DOCUMENT_IMPORTER_DIR),)
DEFAULT_OUT_DIR = Path(output_subdir("06_DocumentImporter"))

# Top-level appendix headers on a PDF page (not "Appendix A-28" catalogue refs)
APPENDIX_HEADER_RE = re.compile(
    r"(?i)^\s*(?:(\d{1,2})\.?\s+)?"
    r"(APPENDIX\s+([A-Z])\b(?!-)"
    r"(?:\s*[-–:]\s*|\s+)(.+?))\s*$"
)
APPENDIX_HEADER_BARE_RE = re.compile(
    r"(?i)^\s*(?:(\d{1,2})\.?\s+)?(APPENDIX\s+([A-Z]))\s*$"
)
SECTION_NUM_ONLY_RE = re.compile(r"^\s*(\d{1,2})\.\s*$")
# Leaf headers inside an appendix (annexes / catalogue items)
ANNEX_LEAF_RE = re.compile(
    r"(?i)^\s*(\d+\.\d+)\s+(App(?:endix)?\s+([A-Z])\s+Annex\b.+?)\s*$"
)
ANNEX_NUM_ONLY_RE = re.compile(r"^\s*(\d+\.\d+)\s*$")
ANNEX_TITLE_RE = re.compile(
    r"(?i)^\s*(App(?:endix)?\s+([A-Z])\s+Annex\b.+?)\s*$"
)
CATALOGUE_ITEM_RE = re.compile(r"^\s*(\d{1,3})\.\s+(.+?)\s*$")
CATALOGUE_NUM_ONLY_RE = re.compile(r"^\s*(\d{1,3})\.\s*$")

SECTION1_HEADING = "Table Of Contents"
SECTION1_BODY = (
    "This section is inserted to align with the section numbering of this "
    "submission submitted by Shui On."
)

# Styles treated as headings in the source document
HEADING_STYLE_LEVEL = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    # Body sections in some ITS docs (e.g. 1701); appendix catalogue items
    # are remapped to level 2 in collect_blocks when region == "appendix"
    "Kaba Heading 1": 1,
    "Kaba Heading 2": 2,
    "Kaba Heading 3": 3,
    "MTRC 1263 Heading 1": 1,  # Appendix B–F titles
}

# Word TOC field / TOC heading styles (not body content)
TOC_FIELD_STYLES = {
    "TOC Heading",
    "toc 1",
    "toc 2",
    "toc 3",
    "toc 4",
    "TOC 1",
    "TOC 2",
    "TOC 3",
    "TOC 4",
}

ACRONYMS = {
    "sams": "SAMS",
    "gcu": "GCU",
    "lcu": "LCU",
    "poe": "POE",
    "ups": "UPS",
    "mcs": "MCS",
    "ssc": "SSC",
    "rt": "RT",
    "sds": "SDS",
    "as": "AS",
    "ibp": "IBP",
    "emi": "EMI",
    "mcb": "MCB",
    "osdp": "OSDP",
    "pin": "PIN",
    "led": "LED",
    "hs": "HS",
    "ms": "MS",
    "lan": "LAN",
    "gi": "G.I.",
    "g.i.": "G.I.",
    "dc": "DC",
    "ac": "AC",
    "rj45": "RJ45",
    "cat6": "CAT6",
    "lszh": "LSZH",
    "rs485": "RS485",
    "sfp": "SFP",
    "abs": "ABS",
    "pc": "PC",
    "uv": "UV",
    "ik08": "IK08",
    "rosh": "RoSH",
    "emf": "EMF",
    "rfi": "RFI",
    "emc": "EMC",
    "ccg": "CCG",
    "amd": "AMD",
    "g2": "G2",
}

LIST_MARKER_RE = re.compile(
    r"^(?:"
    r"[-•\u2022]\s+"
    r"|\([a-z]\)\s+"
    r"|\(\d+\)\s+"
    r"|\d+\.\s+"
    r")",
    re.IGNORECASE,
)

# Prefer "1. Title" / "3.1. Title"; also allow "1 Purpose" (no trailing dot)
PDF_HEADING_RE = re.compile(
    r"^\s*(\d+(?:\.\d+)*)\.\s+(.+?)\s*$"
)
PDF_HEADING_NODOT_RE = re.compile(
    r"^\s*(\d{1,2})\s+([A-Za-z].+?)\s*$"
)
PDF_LIST_RE = re.compile(
    r"^\s*(?:[-•\u2022]|\([a-z]\)|\d+\.)\s+(.+)$",
    re.IGNORECASE,
)


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def normalize_match_key(text: str) -> str:
    """Normalize text for PDF↔DOCX list/heading alignment."""
    t = normalize_spaces(text)
    t = LIST_MARKER_RE.sub("", t)
    t = t.replace("–", "-").replace("—", "-").replace("’", "'").replace("“", '"').replace("”", '"')
    t = re.sub(r"\s+", " ", t).strip().lower()
    return t


def is_appendix_heading(text: str) -> bool:
    t = normalize_spaces(text)
    return bool(re.match(r"(?i)^appendix\s+[a-z]\b", t))


def is_appendix_a(text: str) -> bool:
    t = normalize_spaces(text)
    return bool(re.match(r"(?i)^appendix\s+a\b", t))


def is_table_of_contents(text: str) -> bool:
    return normalize_spaces(text).lower() == "table of contents"


def is_purpose_and_scope(text: str) -> bool:
    return normalize_spaces(text).lower() == "purpose and scope"


def looks_like_section_title(title: str) -> bool:
    """Filter PDF lines that look like body lists/captions, not section titles."""
    t = normalize_spaces(title)
    if len(t) < 2 or len(t) > 80:
        return False
    if t.endswith((".", ";", ",")):
        return False
    if t.count(",") >= 2:
        return False
    if re.match(
        r"(?i)^(the|a|an|this|these|those|it|for|with|and|or|pair|main|welpro)\b",
        t,
    ):
        return False
    # Catalogue / model lines often include digits or quoted vendor names
    if re.search(r"\d{2,}", t):
        return False
    if t.startswith(('"', "'", "“", "‘")):
        return False
    return True


def is_toc_field_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    if style_name in TOC_FIELD_STYLES:
        return True
    return bool(re.match(r"(?i)^toc(\s+\d+)?$", style_name))


_LEAD_TRAIL = (
    r"^([(\"'\u201c\u201d\u2018\u2019\[]*)"
    r"(.*?)"
    r"([)\]\"'\u201c\u201d\u2018\u2019,.:;!?]*)$"
)


def title_case_word(word: str) -> str:
    if not word:
        return word

    # Preserve leading/trailing punctuation around the core token (incl. smart quotes)
    m = re.match(_LEAD_TRAIL, word)
    if not m:
        return word
    lead, core, trail = m.group(1) or "", m.group(2) or "", m.group(3) or ""
    if not core:
        return word

    # Product / model codes: keep original casing when digits are present
    if re.search(r"\d", core):
        return lead + core + trail

    # Hyphenated / slash-separated compounds (no digits)
    if "-" in core or "/" in core:
        sep = "-" if "-" in core else "/"
        parts = core.split(sep)
        return lead + sep.join(title_case_word(p) for p in parts) + trail

    lower = core.lower()
    if lower in ACRONYMS:
        return lead + ACRONYMS[lower] + trail

    # Keep short all-caps tokens that look like acronyms (SAMS, GCU, LCU, …)
    if core.isupper() and 2 <= len(core) <= 4 and core.isalpha():
        return lead + core + trail

    # Keep mixed-case tokens that already look intentional (e.g. PowerEdge)
    if any(c.islower() for c in core[1:]) and any(c.isupper() for c in core[1:]):
        return lead + core[0].upper() + core[1:] + trail

    return lead + core[0].upper() + core[1:].lower() + trail


def to_title_case(text: str) -> str:
    text = normalize_spaces(text)
    if not text:
        return text
    # Normalize dashes to en-dash
    text = text.replace("—", "–")
    words = text.split(" ")
    return " ".join(title_case_word(w) for w in words)


def heading_level(style_name: str | None, text: str) -> int | None:
    if style_name and style_name in HEADING_STYLE_LEVEL:
        level = HEADING_STYLE_LEVEL[style_name]
        # Appendix titles always Heading 1 even if style varies
        if is_appendix_heading(text):
            return 1
        return level
    return None


def paragraph_has_num_pr(paragraph: Paragraph) -> bool:
    return paragraph._p.find(".//" + qn("w:numPr")) is not None


def strip_list_marker(text: str) -> str:
    return LIST_MARKER_RE.sub("", normalize_spaces(text)).strip()


def format_brackets(parts: tuple[int, ...]) -> str:
    return "(" + ".".join(str(p) for p in parts) + ")"


# ---------------------------------------------------------------------------
# PDF outline / list cues (source of truth for structure & list membership)
# ---------------------------------------------------------------------------


@dataclass
class PdfOutline:
    """Heading number map and list-item text keys from PDF layout text."""

    heading_by_title: dict[str, tuple[int, ...]] = field(default_factory=dict)
    list_keys: set[str] = field(default_factory=set)
    first_section_is_toc: bool = False

    def heading_parts(self, title: str) -> tuple[int, ...] | None:
        key = normalize_match_key(title)
        if key in self.heading_by_title:
            return self.heading_by_title[key]
        # Relaxed: prefix match for truncated PDF lines
        for k, parts in self.heading_by_title.items():
            if k.startswith(key) or key.startswith(k):
                if min(len(k), len(key)) >= 12:
                    return parts
        return None

    def is_list_text(self, text: str) -> bool:
        """Exact key membership (used as optional cue; prefer Word numPr)."""
        key = normalize_match_key(text)
        return bool(key) and key in self.list_keys


def _pdf_body_start(raw: str) -> int:
    """Index to begin outline parsing (skip cover / revision noise when possible)."""
    upper = raw.upper()
    markers = (
        "TABLE OF CONTENTS",
        "PURPOSE AND SCOPE",
        "1.   INTRODUCTION",
        "1. INTRODUCTION",
        "1 INTRODUCTION",
        "1.   PURPOSE",
        "1. PURPOSE",
        "1 PURPOSE",
    )
    positions = [upper.find(m) for m in markers if upper.find(m) >= 0]
    return min(positions) if positions else 0


def load_pdf_outline(pdf_path: Path) -> PdfOutline:
    outline = PdfOutline()
    if not pdf_path.exists():
        print(f"Warning: PDF not found ({pdf_path}); list cues limited to Word numPr")
        return outline

    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        print("Warning: pdftotext not on PATH; list cues limited to Word numPr")
        return outline

    raw = subprocess.check_output(
        [pdftotext, "-layout", str(pdf_path), "-"],
        text=True,
        encoding="utf-8",
        errors="replace",
    )

    body = raw[_pdf_body_start(raw) :]
    first_top_level: tuple[str, tuple[int, ...]] | None = None

    for line in body.splitlines():
        stripped_leaders = re.sub(r"\s*\.{2,}.*$", "", line).strip()
        # Skip dotted TOC leaders, but still capture "1. TABLE OF CONTENTS........ N"
        if "...." in line or re.search(r"\.{3,}\s*\d+\s*$", line):
            hm_toc = PDF_HEADING_RE.match(stripped_leaders) or PDF_HEADING_NODOT_RE.match(
                stripped_leaders
            )
            if hm_toc and is_table_of_contents(hm_toc.group(2)):
                parts = tuple(int(x) for x in hm_toc.group(1).split("."))
                key = normalize_match_key(hm_toc.group(2))
                if key and key not in outline.heading_by_title:
                    outline.heading_by_title[key] = parts
                    if first_top_level is None and len(parts) == 1:
                        first_top_level = (key, parts)
                if parts == (1,):
                    outline.first_section_is_toc = True
            continue
        # Page footers
        if re.match(r"^\s*\d+\s*$", line):
            continue
        if "Final Equipment Design Specification" in line:
            continue
        if line.strip().startswith("Submission No."):
            continue
        if re.search(r"(?i)\bPage\s+\d+\s+of\s+\d+\b", line):
            continue

        hm = PDF_HEADING_RE.match(line) or PDF_HEADING_NODOT_RE.match(line)
        if hm:
            num_s, title = hm.group(1), hm.group(2).strip()
            title = re.sub(r"\s*\.{2,}.*$", "", title).strip()
            # Skip revision-history style short lines that look like "1. Added …"
            if re.match(r"(?i)^(added|replaced|updated|removed|revised)\b", title):
                continue
            if not title:
                continue
            parts = tuple(int(x) for x in num_s.split("."))
            if parts[0] > 40:
                continue
            is_toc = is_table_of_contents(title)
            # Section 1: TOC always; other section-1 titles only if title-like.
            # Section 2+: same title-like filter (avoids list/caption noise).
            if not is_toc and not looks_like_section_title(title):
                continue
            key = normalize_match_key(title)
            if key and key not in outline.heading_by_title:
                outline.heading_by_title[key] = parts
                if first_top_level is None and len(parts) == 1:
                    first_top_level = (key, parts)
                if is_toc and parts == (1,):
                    outline.first_section_is_toc = True
            continue

        lm = PDF_LIST_RE.match(line)
        if lm:
            item = lm.group(1).strip()
            # Skip TOC entries mistaken as lists
            if "...." in item or re.search(r"\.{3,}", item):
                continue
            key = normalize_match_key(item)
            if len(key) >= 8:
                outline.list_keys.add(key)
            continue

        # Bullets that PDF encodes with leading spaces + dash already covered;
        # also catch lines that start with (a)/(b) without the generic regex miss
        am = re.match(r"^\s*\(([a-z])\)\s+(.+)$", line, re.I)
        if am:
            key = normalize_match_key(am.group(2))
            if len(key) >= 8:
                outline.list_keys.add(key)

    if first_top_level and is_table_of_contents(first_top_level[0]):
        outline.first_section_is_toc = True
    else:
        toc_parts = outline.heading_parts("Table Of Contents")
        outline.first_section_is_toc = toc_parts == (1,)

    print(
        f"PDF outline: {len(outline.heading_by_title)} headings, "
        f"{len(outline.list_keys)} list cues, "
        f"first_section_is_toc={outline.first_section_is_toc}"
    )
    return outline


# ---------------------------------------------------------------------------
# Document model + numbering
# ---------------------------------------------------------------------------


@dataclass
class Block:
    kind: str  # heading | para | table
    level: int | None = None
    text: str = ""
    titled: str = ""
    is_list: bool = False
    has_drawing: bool = False
    paragraph: Paragraph | None = None
    table: Table | None = None
    region: str = "body"  # body | appendix
    path: tuple[int, ...] | None = None
    bracket_path: tuple[int, ...] | None = None
    local_num: int | None = None


def collect_blocks(
    src: Document,
    pdf: PdfOutline,
    *,
    appendix_headers_only: bool = False,
) -> list[Block]:
    """Collect body blocks starting at the first real Heading 1 (skip TOC fields).

    When appendix_headers_only is True, appendix regions keep headings only
    (FEDS-style); paragraphs and tables under appendices are skipped.
    """
    blocks: list[Block] = []
    mode = "skip"  # skip | body
    in_appendix = False

    for kind, raw in iter_body_blocks(src):
        if kind == "p":
            paragraph: Paragraph = raw
            text = paragraph.text or ""
            sn = paragraph.style.name if paragraph.style else ""
            if is_toc_field_style(sn):
                continue
            level = heading_level(sn, text)
            norm = normalize_spaces(text)

            if mode == "skip":
                if level == 1 and norm:
                    if is_table_of_contents(norm):
                        # TOC is section 1 in DOCX; content starts at the next H1
                        continue
                    mode = "body"
                else:
                    continue

            assert mode == "body"
            if level == 1 and norm and is_appendix_heading(norm):
                in_appendix = True
            region = "appendix" if in_appendix else "body"

            if level and norm:
                out_level = level
                if region == "appendix":
                    out_level = 1 if is_appendix_heading(norm) else min(level, 4)
                    if not is_appendix_heading(norm) and sn == "Kaba Heading 1":
                        out_level = 2
                blocks.append(
                    Block(
                        kind="heading",
                        level=out_level,
                        text=norm,
                        titled=to_title_case(norm),
                        paragraph=paragraph,
                        region=region,
                    )
                )
            elif level and not norm:
                continue
            else:
                if appendix_headers_only and region == "appendix":
                    continue
                if not norm and not paragraph_has_drawing(paragraph._p):
                    continue
                # Word list markup is authoritative for DOCX extraction.
                is_list = False
                if norm:
                    is_list = paragraph_has_num_pr(paragraph) or (
                        sn == "List Paragraph"
                    )
                blocks.append(
                    Block(
                        kind="para",
                        text=norm,
                        is_list=is_list,
                        has_drawing=paragraph_has_drawing(paragraph._p),
                        paragraph=paragraph,
                        region=region,
                    )
                )

        elif kind == "tbl":
            if mode != "body":
                continue
            if appendix_headers_only and in_appendix:
                continue
            blocks.append(
                Block(
                    kind="table",
                    table=raw,
                    region="appendix" if in_appendix else "body",
                )
            )

    return blocks


def document_begins_with_toc(pdf: PdfOutline, blocks: list[Block], src: Document) -> bool:
    """True when the document's first numbered/body section is Table Of Contents."""
    if pdf.first_section_is_toc:
        return True
    # DOCX: first Heading 1 in document order (including skipped TOC heading)
    for kind, raw in iter_body_blocks(src):
        if kind != "p":
            continue
        paragraph: Paragraph = raw
        sn = paragraph.style.name if paragraph.style else ""
        if is_toc_field_style(sn):
            # Unnumbered TOC page does not count as numbered section 1
            continue
        level = heading_level(sn, paragraph.text or "")
        norm = normalize_spaces(paragraph.text or "")
        if level == 1 and norm:
            return is_table_of_contents(norm)
    return False


def assign_heading_paths(blocks: list[Block], pdf: PdfOutline) -> None:
    """Assign hierarchical paths to headings (PDF map when available)."""
    counters = [0, 0, 0, 0]  # levels 1..4
    for b in blocks:
        if b.kind != "heading":
            continue
        L = b.level or 1
        pdf_parts = pdf.heading_parts(b.text)
        if pdf_parts and len(pdf_parts) == L:
            # Sync counters to PDF numbering
            for i, v in enumerate(pdf_parts):
                counters[i] = v
            for i in range(L, 4):
                counters[i] = 0
            b.path = pdf_parts
        else:
            counters[L - 1] += 1
            for i in range(L, 4):
                counters[i] = 0
            b.path = tuple(counters[i] for i in range(L))


def assign_body_brackets(blocks: list[Block]) -> None:
    """Assign bracket paths to body paragraphs (incl. .0 gaps)."""
    for i, b in enumerate(blocks):
        if b.kind != "para" or not b.text:
            continue

        parent_idx = None
        parent: Block | None = None
        for j in range(i - 1, -1, -1):
            if blocks[j].kind == "heading" and blocks[j].path:
                parent_idx = j
                parent = blocks[j]
                break
        if parent is None or parent.path is None or parent_idx is None:
            continue

        L = len(parent.path)

        section_end = len(blocks)
        for j in range(parent_idx + 1, len(blocks)):
            bj = blocks[j]
            if bj.kind == "heading" and bj.level is not None and bj.level <= L:
                section_end = j
                break

        child_indices = [
            j
            for j in range(parent_idx + 1, section_end)
            if blocks[j].kind == "heading"
            and blocks[j].level is not None
            and blocks[j].level == L + 1
        ]

        if child_indices:
            first_child = child_indices[0]
            if i < first_child:
                bucket = [
                    j
                    for j in range(parent_idx + 1, first_child)
                    if blocks[j].kind == "para" and blocks[j].text
                ]
                n = bucket.index(i) + 1
                b.bracket_path = parent.path + (0, n)
            # else: nearest-parent logic should have selected a deeper heading
        else:
            bucket = [
                j
                for j in range(parent_idx + 1, section_end)
                if blocks[j].kind == "para" and blocks[j].text
            ]
            n = bucket.index(i) + 1
            b.bracket_path = parent.path + (n,)


def assign_local_list_numbers(blocks: list[Block]) -> None:
    prev_list = False
    local_i = 0
    for b in blocks:
        if b.kind == "para" and b.is_list and b.text:
            if not prev_list:
                local_i = 1
            else:
                local_i += 1
            b.local_num = local_i
            prev_list = True
        else:
            # headings, tables, non-list paras, image-only paras reset the list
            if b.kind in ("heading", "table") or (
                b.kind == "para" and (not b.is_list or not b.text)
            ):
                prev_list = False


def numbered_body_text(b: Block) -> str:
    """Build display text: (x.y.z) [N. ]stripped_text"""
    core = strip_list_marker(b.text) if b.is_list else normalize_spaces(b.text)
    if b.local_num is not None:
        core = f"{b.local_num}. {core}"
    if b.bracket_path:
        return f"{format_brackets(b.bracket_path)} {core}"
    return core


def numbered_heading_text(b: Block) -> str:
    title = b.titled or to_title_case(b.text)
    if b.path:
        return f"{format_brackets(b.path)} {title}"
    return title


# ---------------------------------------------------------------------------
# DOCX / image helpers
# ---------------------------------------------------------------------------


def iter_body_blocks(doc: Document):
    """Yield ('p', Paragraph) or ('tbl', Table) in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "tbl", Table(child, doc)


def paragraph_has_drawing(element) -> bool:
    xml = element.xml
    return "w:drawing" in xml or "v:imagedata" in xml or "w:pict" in xml


def table_has_drawing(table: Table) -> bool:
    return paragraph_has_drawing(table._tbl)


def collect_embed_rids(element) -> list[str]:
    rids: list[str] = []
    for blip in element.findall(".//" + A_BLIP):
        rid = blip.get(R_EMBED)
        if rid:
            rids.append(rid)
    for imagedata in element.findall(".//" + VML_IMAGEDATA):
        rid = imagedata.get(R_ID)
        if rid:
            rids.append(rid)
    return rids


class ImageStore:
    """Extract images from source and remap into destination DOCX + media folder."""

    def __init__(self, src_doc: Document, media_dir: Path):
        self.src_doc = src_doc
        self.media_dir = media_dir
        self.media_dir.mkdir(parents=True, exist_ok=True)
        self._hash_to_file: dict[str, Path] = {}
        self._hash_to_dst_rid: dict[str, str] = {}
        self._src_rid_to_hash: dict[str, str] = {}
        self._img_index = 0
        self._diagram_drawing_copied_for: set[int] = set()

    def _part_for_rid(self, rid: str):
        rel = self.src_doc.part.rels.get(rid)
        if rel is None:
            return None
        if rel.is_external:
            return None
        return rel.target_part

    def ensure_extracted(self, rid: str) -> Path | None:
        if rid in self._src_rid_to_hash:
            h = self._src_rid_to_hash[rid]
            return self._hash_to_file[h]
        part = self._part_for_rid(rid)
        if part is None:
            return None
        blob = part.blob
        h = hashlib.sha1(blob).hexdigest()
        self._src_rid_to_hash[rid] = h
        if h in self._hash_to_file:
            return self._hash_to_file[h]
        self._img_index += 1
        ext = Path(str(part.partname)).suffix.lstrip(".") or "bin"
        filename = f"image_{self._img_index:04d}.{ext}"
        path = self.media_dir / filename
        path.write_bytes(blob)
        self._hash_to_file[h] = path
        return path

    def remap_element_images(self, dst_doc: Document, element) -> None:
        """Rewrite image and SmartArt rIds on a deepcopy'd element for dst_doc."""
        for blip in element.findall(".//" + A_BLIP):
            old = blip.get(R_EMBED)
            if not old:
                continue
            new_rid = self._ensure_dst_rid(dst_doc, old)
            if new_rid:
                blip.set(R_EMBED, new_rid)
            else:
                blip.attrib.pop(R_EMBED, None)
        for imagedata in element.findall(".//" + VML_IMAGEDATA):
            old = imagedata.get(R_ID)
            if not old:
                continue
            new_rid = self._ensure_dst_rid(dst_doc, old)
            if new_rid:
                imagedata.set(R_ID, new_rid)
        self._remap_diagrams(dst_doc, element)

    def _remap_diagrams(self, dst_doc: Document, element) -> None:
        """Copy SmartArt diagram parts and rewrite dgm:relIds, or strip on failure."""
        for rel_ids in list(element.findall(".//" + DGM_REL_IDS)):
            new_attrs: dict[str, str] = {}
            ok = True
            for attr, rt in DGM_ATTR_TO_RT.items():
                key = f"{{{R_NS}}}{attr}"
                old = rel_ids.get(key)
                if not old:
                    ok = False
                    break
                partname_tmpl = DGM_RT_TO_PARTNAME[rt]
                new_rid = self._ensure_dst_part_rid(dst_doc, old, rt, partname_tmpl)
                if not new_rid:
                    ok = False
                    break
                new_attrs[key] = new_rid
            if ok:
                for key, new_rid in new_attrs.items():
                    rel_ids.set(key, new_rid)
                self._ensure_diagram_drawing_parts(dst_doc)
            else:
                self._replace_diagram_with_placeholder(rel_ids)

    def _ensure_diagram_drawing_parts(self, dst_doc: Document) -> None:
        """Copy MS diagramDrawing parts from source (often present alongside SmartArt)."""
        dst_id = id(dst_doc.part)
        if dst_id in self._diagram_drawing_copied_for:
            return
        self._diagram_drawing_copied_for.add(dst_id)
        for rel in self.src_doc.part.rels.values():
            if rel.is_external or rel.reltype != RT_DIAGRAM_DRAWING:
                continue
            self._ensure_dst_part_rid(
                dst_doc,
                rel.rId,
                RT_DIAGRAM_DRAWING,
                DGM_RT_TO_PARTNAME[RT_DIAGRAM_DRAWING],
            )

    @staticmethod
    def _replace_diagram_with_placeholder(rel_ids_el) -> None:
        """Remove a broken SmartArt drawing and leave placeholder text in its run."""
        drawing = rel_ids_el
        while drawing is not None and drawing.tag != W_DRAWING:
            drawing = drawing.getparent()
        if drawing is None:
            parent = rel_ids_el.getparent()
            if parent is not None:
                parent.remove(rel_ids_el)
            return
        run = drawing.getparent()
        run.remove(drawing)
        if run is not None and run.tag == W_R:
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = SMARTART_PLACEHOLDER
            run.append(t)

    def _ensure_dst_rid(self, dst_doc: Document, src_rid: str) -> str | None:
        path = self.ensure_extracted(src_rid)
        if path is None:
            return None
        part = self._part_for_rid(src_rid)
        if part is None:
            return None
        h = self._src_rid_to_hash[src_rid]
        cache_key = f"{id(dst_doc.part)}:{h}"
        if cache_key in self._hash_to_dst_rid:
            return self._hash_to_dst_rid[cache_key]

        blob = part.blob
        content_type = part.content_type
        ext = Path(str(part.partname)).suffix.lstrip(".") or "bin"
        package = dst_doc.part.package
        partname = package.next_partname(PackURI(f"/word/media/image%d.{ext}"))
        image_part = ImagePart(partname, content_type, blob, package)
        new_rid = dst_doc.part.relate_to(image_part, RT.IMAGE)
        self._hash_to_dst_rid[cache_key] = new_rid
        return new_rid

    def _ensure_dst_part_rid(
        self,
        dst_doc: Document,
        src_rid: str,
        reltype: str,
        partname_tmpl: str,
    ) -> str | None:
        """Copy an arbitrary source part (e.g. SmartArt) into dst_doc and relate it."""
        part = self._part_for_rid(src_rid)
        if part is None:
            return None
        h = hashlib.sha1(part.blob).hexdigest()
        cache_key = f"{id(dst_doc.part)}:{reltype}:{h}"
        if cache_key in self._hash_to_dst_rid:
            return self._hash_to_dst_rid[cache_key]
        package = dst_doc.part.package
        partname = package.next_partname(PackURI(partname_tmpl))
        new_part = Part(partname, part.content_type, part.blob, package)
        new_rid = dst_doc.part.relate_to(new_part, reltype)
        self._hash_to_dst_rid[cache_key] = new_rid
        return new_rid


def assert_document_rels_intact(doc: Document) -> None:
    """Raise if document.xml references rIds missing from document.xml.rels."""
    used: set[str] = set()
    r_ns_prefix = "{" + R_NS + "}"
    for el in doc.part.element.iter():
        for key, val in el.attrib.items():
            if not isinstance(val, str) or not val.startswith("rId"):
                continue
            if key.startswith(r_ns_prefix) or key.startswith("r:"):
                used.add(val)
    defined = set(doc.part.rels.keys())
    missing = sorted(used - defined)
    if missing:
        raise ValueError(
            "Refusing to save DOCX with orphaned relationship IDs: "
            + ", ".join(missing)
        )


def _set_run_font(run: Run) -> None:
    """Force Arial 12pt on a run, including East Asian / complex-script slots."""
    run.font.name = OUTPUT_FONT_NAME
    run.font.size = OUTPUT_FONT_SIZE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), OUTPUT_FONT_NAME)
    rFonts.set(qn("w:hAnsi"), OUTPUT_FONT_NAME)
    rFonts.set(qn("w:eastAsia"), OUTPUT_FONT_NAME)
    rFonts.set(qn("w:cs"), OUTPUT_FONT_NAME)


def _set_style_font(style) -> None:
    style.font.name = OUTPUT_FONT_NAME
    style.font.size = OUTPUT_FONT_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), OUTPUT_FONT_NAME)
    rFonts.set(qn("w:hAnsi"), OUTPUT_FONT_NAME)
    rFonts.set(qn("w:eastAsia"), OUTPUT_FONT_NAME)
    rFonts.set(qn("w:cs"), OUTPUT_FONT_NAME)


def apply_output_fonts(doc: Document) -> None:
    """Set Normal + Heading 1-4 styles to Arial 12pt."""
    _set_style_font(doc.styles["Normal"])
    for level in range(1, 5):
        try:
            _set_style_font(doc.styles[f"Heading {level}"])
        except KeyError:
            pass


def force_paragraph_font(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        _set_run_font(run)


def force_table_font(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                force_paragraph_font(paragraph)


def _make_border_el(tag: str):
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), OUTPUT_TABLE_BORDER_SZ)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), OUTPUT_TABLE_BORDER_COLOR)
    return el


def force_table_borders(table: Table) -> None:
    """Force visible black grid borders on the table (and each cell).

    Copied tables often rely on styles (e.g. GridTable1Light) that are absent
    from a blank Document, so borders would otherwise disappear or stay 'auto'.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_make_border_el(edge))
    tblPr.append(borders)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            existing_tc = tcPr.find(qn("w:tcBorders"))
            if existing_tc is not None:
                tcPr.remove(existing_tc)
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                tc_borders.append(_make_border_el(edge))
            tcPr.append(tc_borders)


def set_run_text(paragraph: Paragraph, text: str, *, bold: bool | None = None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    _set_run_font(run)


def add_heading_para(doc: Document, text: str, level: int) -> Paragraph:
    style = f"Heading {level}"
    para = doc.add_paragraph(text, style=style)
    force_paragraph_font(para)
    return para


def add_body_para(doc: Document, text: str, style: str = "Normal") -> Paragraph:
    try:
        para = doc.add_paragraph(text, style=style)
    except KeyError:
        para = doc.add_paragraph(text, style="Normal")
    force_paragraph_font(para)
    return para


def copy_paragraph_to_doc(
    src_doc: Document,
    dst_doc: Document,
    paragraph: Paragraph,
    images: ImageStore,
    *,
    force_style: str | None = None,
    force_text: str | None = None,
) -> Paragraph:
    new_p = dst_doc.add_paragraph()
    if force_text is not None:
        style_name = force_style or "Normal"
        try:
            new_p.style = style_name
        except KeyError:
            new_p.style = "Normal"
        set_run_text(new_p, force_text)
        return new_p

    new_el = deepcopy(paragraph._p)
    # Always strip Word auto-numbering so lists cannot continue
    pPr = new_el.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)

    images.remap_element_images(dst_doc, new_el)
    new_p._p.getparent().replace(new_p._p, new_el)
    new_p = Paragraph(new_el, dst_doc)
    if force_style:
        try:
            new_p.style = force_style
        except KeyError:
            pass
    else:
        try:
            new_p.style = "Normal"
        except KeyError:
            pass
    force_paragraph_font(new_p)
    return new_p


def copy_table_to_doc(
    src_doc: Document,
    dst_doc: Document,
    table: Table,
    images: ImageStore,
) -> Table:
    new_tbl_el = deepcopy(table._tbl)
    images.remap_element_images(dst_doc, new_tbl_el)
    body = dst_doc.element.body
    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(new_tbl_el)
    else:
        body.append(new_tbl_el)
    new_table = Table(new_tbl_el, dst_doc)
    force_table_font(new_table)
    force_table_borders(new_table)
    return new_table


def table_to_markdown(table: Table, images: ImageStore, media_rel: str) -> str:
    if table_has_drawing(table):
        rows_out = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = normalize_spaces(cell.text)
                has_img = "w:drawing" in cell._tc.xml or "v:imagedata" in cell._tc.xml
                if has_img:
                    rids = collect_embed_rids(cell._tc)
                    imgs = []
                    for rid in rids:
                        path = images.ensure_extracted(rid)
                        if path:
                            imgs.append(f"![]({media_rel}/{path.name})")
                    if text:
                        cells.append(text + " " + " ".join(imgs))
                    else:
                        cells.append(" ".join(imgs) if imgs else "[image]")
                else:
                    cells.append(text.replace("|", "\\|"))
            rows_out.append(cells)
        if not rows_out:
            return "<!-- TABLE: empty, preserved in DOCX -->\n"
        width = max(len(r) for r in rows_out)
        lines = ["<!-- TABLE with images: layout preserved in DOCX -->"]
        header = rows_out[0] + [""] * (width - len(rows_out[0]))
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in rows_out[1:]:
            padded = row + [""] * (width - len(row))
            lines.append("| " + " | ".join(padded) + " |")
        return "\n".join(lines) + "\n"

    rows_out = []
    for row in table.rows:
        cells = [normalize_spaces(c.text).replace("|", "\\|") for c in row.cells]
        rows_out.append(cells)
    if not rows_out:
        return ""
    width = max(len(r) for r in rows_out)
    lines = []
    header = rows_out[0] + [""] * (width - len(rows_out[0]))
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in rows_out[1:]:
        padded = row + [""] * (width - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines) + "\n"


def image_para_to_markdown(paragraph: Paragraph, images: ImageStore, media_rel: str) -> str:
    rids = collect_embed_rids(paragraph._p)
    img_md = []
    for rid in rids:
        path = images.ensure_extracted(rid)
        if path:
            img_md.append(f"![]({media_rel}/{path.name})")
    if img_md:
        return "\n\n".join(img_md) + "\n"
    return ""


# ---------------------------------------------------------------------------
# Multi-page appendix PDF extraction
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class AppendixSpan:
    letter: str
    title: str
    start_page: int  # 0-based inclusive
    end_page: int  # 0-based inclusive

    @property
    def page_count(self) -> int:
        return self.end_page - self.start_page + 1


@dataclass(frozen=True)
class AppendixLeafSpan:
    letter: str
    index_label: str  # e.g. "01" or "8.1"
    title: str
    start_page: int  # 0-based inclusive
    end_page: int  # 0-based inclusive

    @property
    def page_count(self) -> int:
        return self.end_page - self.start_page + 1


def _page_is_toc_heavy(text: str) -> bool:
    dotted = len(re.findall(r"\.{3,}", text))
    return dotted >= 5


def _match_appendix_header(
    line: str, prev_line: str | None = None
) -> tuple[str, str, bool] | None:
    """Return (letter, title, has_section_number) if line is an appendix header."""
    m = APPENDIX_HEADER_RE.match(line) or APPENDIX_HEADER_BARE_RE.match(line)
    if not m:
        return None
    sec, full, letter = m.group(1), m.group(2), m.group(3).upper()
    title = normalize_spaces(full)
    # Cross-doc IRS refs, e.g. "Appendix I Annex 27" — not this document's appendix
    if re.search(r"(?i)^APPENDIX\s+[A-Z]\s+ANNEX\b", title):
        return None
    # Trailing-period short refs in body text, e.g. "Appendix A - I/O List."
    rest = m.group(4) if m.lastindex and m.lastindex >= 4 else ""
    if rest is None:
        rest = ""
    if not sec and rest.strip().endswith(".") and len(rest.strip()) <= 40:
        return None
    has_sec = bool(sec) or bool(
        prev_line and SECTION_NUM_ONLY_RE.match(prev_line or "")
    )
    return letter, title, has_sec


def _match_annex_leaf(
    line: str, appendix_letter: str, next_line: str | None = None
) -> tuple[str, str] | None:
    """Return (index_label, title) for App X Annex headers inside an appendix."""
    m = ANNEX_LEAF_RE.match(line)
    if m:
        if m.group(3).upper() != appendix_letter.upper():
            return None
        return m.group(1), normalize_spaces(m.group(2))

    # Split across lines: "10.1" then "App C Annex 1 - …"
    m_num = ANNEX_NUM_ONLY_RE.match(line)
    if m_num and next_line:
        m_title = ANNEX_TITLE_RE.match(next_line)
        if not m_title:
            return None
        if m_title.group(2).upper() != appendix_letter.upper():
            return None
        return m_num.group(1), normalize_spaces(m_title.group(1))
    return None


def _match_catalogue_leaf(
    line: str, next_line: str | None = None
) -> tuple[str, str] | None:
    """Return (index_label, title) for numbered catalogue item headers."""
    m = CATALOGUE_ITEM_RE.match(line)
    if m:
        num, rest = m.group(1), normalize_spaces(m.group(2))
        if re.match(r"(?i)^APPENDIX\b", rest):
            return None
        if len(rest) < 3:
            return None
        return num, rest

    m_num = CATALOGUE_NUM_ONLY_RE.match(line)
    if m_num and next_line:
        rest = normalize_spaces(next_line)
        if re.match(r"(?i)^APPENDIX\b", rest):
            return None
        if len(rest) < 3:
            return None
        # Avoid treating section numbers before appendix titles as catalogue items
        if _match_appendix_header(next_line, line):
            return None
        # Bare "1." before drawing notes / dimensions — require title-like text
        if not re.match(r"^[\"'“”A-Za-z]", rest):
            return None
        return m_num.group(1), rest
    return None


def find_appendix_spans(pdf_path: Path) -> list[AppendixSpan]:
    """Locate top-level appendix ranges in a PDF (cover page through next appendix)."""
    doc = fitz.open(pdf_path)
    try:
        # letter -> list of (page_index, title, has_section_number)
        candidates: dict[str, list[tuple[int, str, bool]]] = {}
        for i in range(doc.page_count):
            text = doc.load_page(i).get_text("text")
            if _page_is_toc_heavy(text):
                continue
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            prev: str | None = None
            # Headers appear near the top after running headers/footers
            for ln in lines[:30]:
                hit = _match_appendix_header(ln, prev)
                prev = ln
                if not hit:
                    continue
                letter, title, has_sec = hit
                candidates.setdefault(letter, []).append((i, title, has_sec))
                break

        chosen: list[tuple[str, str, int]] = []
        for letter, hits in candidates.items():
            # Prefer section-numbered headers; then the latest page (body over TOC/refs)
            ranked = sorted(hits, key=lambda h: (h[2], h[0]))
            page_i, title, _ = ranked[-1]
            chosen.append((letter, title, page_i))

        chosen.sort(key=lambda x: x[2])
        spans: list[AppendixSpan] = []
        for idx, (letter, title, start) in enumerate(chosen):
            end = (
                chosen[idx + 1][2] - 1
                if idx + 1 < len(chosen)
                else doc.page_count - 1
            )
            if end < start:
                end = start
            spans.append(
                AppendixSpan(
                    letter=letter,
                    title=title,
                    start_page=start,
                    end_page=end,
                )
            )
        return spans
    finally:
        doc.close()


def _collect_leaf_starts(
    doc: fitz.Document, span: AppendixSpan
) -> list[tuple[int, str, str]]:
    """
    Find leaf header starts inside an appendix span.
    Returns list of (page_index, index_label, title), sorted by page.
    Prefers annex-style headers when present; otherwise catalogue items.
    """
    annex_hits: list[tuple[int, str, str]] = []
    catalogue_hits: list[tuple[int, str, str]] = []

    for i in range(span.start_page, span.end_page + 1):
        text = doc.load_page(i).get_text("text")
        if _page_is_toc_heavy(text):
            continue
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for idx, ln in enumerate(lines[:30]):
            nxt = lines[idx + 1] if idx + 1 < len(lines) else None
            annex = _match_annex_leaf(ln, span.letter, nxt)
            if annex:
                annex_hits.append((i, annex[0], annex[1]))
                break
            cat = _match_catalogue_leaf(ln, nxt)
            if cat:
                catalogue_hits.append((i, cat[0], cat[1]))
                break

    # Annex leaves win when present. Catalogue needs 2+ hits to avoid drawing-note false positives.
    if annex_hits:
        hits = annex_hits
    elif len(catalogue_hits) >= 2:
        hits = catalogue_hits
    else:
        return []

    # Deduplicate: keep first occurrence of each index_label, in page order
    hits.sort(key=lambda h: h[0])
    seen: set[str] = set()
    unique: list[tuple[int, str, str]] = []
    for page_i, label, title in hits:
        if label in seen:
            continue
        seen.add(label)
        unique.append((page_i, label, title))
    return unique


def find_appendix_leaf_spans(
    doc: fitz.Document, span: AppendixSpan
) -> list[AppendixLeafSpan]:
    """Build leaf page ranges under a top-level appendix (empty if none)."""
    starts = _collect_leaf_starts(doc, span)
    if not starts:
        return []

    # Include appendix cover with the first leaf
    first_page, first_label, first_title = starts[0]
    if first_page > span.start_page:
        starts[0] = (span.start_page, first_label, first_title)
    elif first_page < span.start_page:
        starts[0] = (span.start_page, first_label, first_title)

    leaves: list[AppendixLeafSpan] = []
    for idx, (start, label, title) in enumerate(starts):
        end = (
            starts[idx + 1][0] - 1
            if idx + 1 < len(starts)
            else span.end_page
        )
        if end < start:
            end = start
        # Zero-pad plain integer catalogue indexes for stable filenames
        if label.isdigit():
            index_label = f"{int(label):02d}"
        else:
            index_label = label
        leaves.append(
            AppendixLeafSpan(
                letter=span.letter,
                index_label=index_label,
                title=title,
                start_page=start,
                end_page=end,
            )
        )
    return leaves


def _slugify_title(title: str, max_len: int = 80) -> str:
    slug = re.sub(r"[^\w]+", "_", title, flags=re.UNICODE).strip("_")
    slug = re.sub(r"_+", "_", slug)
    if len(slug) > max_len:
        slug = slug[:max_len].rstrip("_")
    return slug


def _appendix_filename(stem: str, span: AppendixSpan) -> str:
    # Slug first so files sort/read by appendix title, then source doc
    return f"{_slugify_title(span.title)}_{stem}.pdf"


def _appendix_leaf_filename(stem: str, leaf: AppendixLeafSpan) -> str:
    slug = _slugify_title(leaf.title, max_len=70)
    return f"Appendix_{leaf.letter}_{leaf.index_label}_{slug}_{stem}.pdf"


def extract_appendix_pdfs(
    pdf_path: Path, out_dir: Path, stem: str | None = None
) -> list[Path]:
    """
    Write separate PDFs for multi-page appendices.

    When an appendix contains leaf headers (annexes / catalogue items), emit
    one PDF per multi-page leaf instead of a single giant appendix PDF.
    """
    stem = stem or pdf_path.stem
    spans = find_appendix_spans(pdf_path)

    app_dir = out_dir / "appendices" / stem
    if app_dir.exists():
        for old in app_dir.glob("*.pdf"):
            old.unlink()
    app_dir.mkdir(parents=True, exist_ok=True)

    src = fitz.open(pdf_path)
    written: list[Path] = []
    try:
        for s in spans:
            # if s.page_count <= 1:
            #     print(
            #         f"  Appendix {s.letter}: pages {s.start_page + 1}-{s.end_page + 1} "
            #         f"({s.page_count} page) — skip (single-page)"
            #     )
            #     continue

            leaves = find_appendix_leaf_spans(src, s)
            if leaves:
                print(
                    f"  Appendix {s.letter}: {len(leaves)} leaf header(s) "
                    f"(pages {s.start_page + 1}-{s.end_page + 1})"
                )
                for leaf in leaves:
                    out_path = app_dir / _appendix_leaf_filename(stem, leaf)
                    dst = fitz.open()
                    try:
                        dst.insert_pdf(
                            src, from_page=leaf.start_page, to_page=leaf.end_page
                        )
                        dst.save(out_path)
                    finally:
                        dst.close()
                    written.append(out_path)
                    print(
                        f"    [{leaf.index_label}] pages "
                        f"{leaf.start_page + 1}-{leaf.end_page + 1} "
                        f"({leaf.page_count} page{'s' if leaf.page_count != 1 else ''}) "
                        f"-> {out_path.name}"
                    )
                continue

            out_path = app_dir / _appendix_filename(stem, s)
            dst = fitz.open()
            try:
                dst.insert_pdf(src, from_page=s.start_page, to_page=s.end_page)
                dst.save(out_path)
            finally:
                dst.close()
            written.append(out_path)
            print(
                f"  Appendix {s.letter}: pages {s.start_page + 1}-{s.end_page + 1} "
                f"({s.page_count} pages) -> {out_path.name}"
            )
    finally:
        src.close()

    if not written:
        print("  No multi-page appendices/leaves to extract")
    else:
        print(f"Wrote {len(written)} appendix PDF(s) under {app_dir}")
    return written


# ---------------------------------------------------------------------------
# Extract
# ---------------------------------------------------------------------------


def extract(
    src_path: Path,
    out_dir: Path,
    pdf_path: Path | None = None,
    *,
    appendix_headers_only: bool = False,
) -> None:
    stem = src_path.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    media_dir = out_dir / "media" / stem
    media_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = pdf_path or src_path.with_suffix(".pdf")
    print(f"Loading PDF outline from {pdf_path} ...")
    pdf = load_pdf_outline(pdf_path)

    print(f"Loading {src_path} ...")
    src = Document(str(src_path))
    images = ImageStore(src, media_dir)

    if appendix_headers_only:
        print("Appendix mode: headers only (skip appendix body content)")
    blocks = collect_blocks(src, pdf, appendix_headers_only=appendix_headers_only)
    insert_toc = document_begins_with_toc(pdf, blocks, src)
    print(f"Hard-insert TOC section: {insert_toc}")

    assign_heading_paths(blocks, pdf)

    # When synthetic TOC is section 1, Purpose (or first body H1 at path 1) must be >= 2
    if insert_toc:
        for b in blocks:
            if b.kind == "heading" and b.level == 1 and is_purpose_and_scope(b.text):
                if b.path == (1,):
                    for bb in blocks:
                        if bb.path:
                            bb.path = (bb.path[0] + 1,) + bb.path[1:]
                        if bb.bracket_path:
                            bb.bracket_path = (bb.bracket_path[0] + 1,) + bb.bracket_path[1:]
                break

    assign_body_brackets(blocks)
    assign_local_list_numbers(blocks)

    dst = Document()
    apply_output_fonts(dst)

    md_parts: list[str] = []
    media_rel = f"media/{stem}"

    if insert_toc:
        add_heading_para(dst, f"(1) {SECTION1_HEADING}", 1)
        add_body_para(dst, f"(1.1) {SECTION1_BODY}", "Normal")
        md_parts.append(f"# (1) {SECTION1_HEADING}\n")
        md_parts.append(f"(1.1) {SECTION1_BODY}\n")

    stats = {"body_paras": 0, "body_tables": 0, "appendix_headings": 0, "lists": 0}

    for b in blocks:
        if b.kind == "heading":
            assert b.paragraph is not None
            text = numbered_heading_text(b)
            level = b.level or 1
            copy_paragraph_to_doc(
                src,
                dst,
                b.paragraph,
                images,
                force_style=f"Heading {level}",
                force_text=text,
            )
            md_parts.append(f"{'#' * level} {text}\n")
            if b.region == "appendix":
                stats["appendix_headings"] += 1
            else:
                stats["body_paras"] += 1

        elif b.kind == "para":
            assert b.paragraph is not None
            if b.text:
                text = numbered_body_text(b)
                if b.local_num is not None:
                    stats["lists"] += 1
                copy_paragraph_to_doc(
                    src,
                    dst,
                    b.paragraph,
                    images,
                    force_style="Normal",
                    force_text=text,
                )
                md_parts.append(text + "\n")
                stats["body_paras"] += 1
            elif b.has_drawing:
                copy_paragraph_to_doc(src, dst, b.paragraph, images, force_style="Normal")
                md = image_para_to_markdown(b.paragraph, images, media_rel)
                if md:
                    md_parts.append(md if md.endswith("\n") else md + "\n")
                stats["body_paras"] += 1

        elif b.kind == "table":
            assert b.table is not None
            copy_table_to_doc(src, dst, b.table, images)
            md_parts.append(table_to_markdown(b.table, images, media_rel))
            md_parts.append("")
            stats["body_tables"] += 1

    md_path = out_dir / f"{stem}.md"
    docx_path = out_dir / f"{stem}.docx"

    md_text = "\n".join(md_parts)
    md_text = re.sub(r"\n{3,}", "\n\n", md_text).strip() + "\n"
    md_path.write_text(md_text, encoding="utf-8")
    assert_document_rels_intact(dst)
    dst.save(str(docx_path))

    print(f"Wrote {md_path}")
    print(f"Wrote {docx_path}")
    print(
        f"Stats: body_paras={stats['body_paras']}, body_tables={stats['body_tables']}, "
        f"appendix_headings={stats['appendix_headings']}, lists={stats['lists']}, "
        f"images={images._img_index}"
    )

    print(f"Extracting multi-page appendices from {pdf_path} ...")
    extract_appendix_pdfs(pdf_path, out_dir, stem=stem)


def discover_pairs(input_dir: Path) -> list[tuple[Path, Path]]:
    """Return (docx, pdf) pairs; warn about legacy .doc and missing PDFs."""
    pairs: list[tuple[Path, Path]] = []
    docx_files = sorted(input_dir.glob("*.docx"))
    for docx in docx_files:
        pdf = docx.with_suffix(".pdf")
        if not pdf.exists():
            print(f"Warning: skipping {docx.name} (no sibling PDF)", file=sys.stderr)
            continue
        pairs.append((docx, pdf))

    for doc in sorted(input_dir.glob("*.doc")):
        if doc.suffix.lower() != ".doc":
            continue
        print(
            f"Warning: skipping {doc.name} (legacy .doc; convert to .docx and re-run)",
            file=sys.stderr,
        )
    return pairs


def discover_orphan_pdfs(input_dir: Path, paired_pdfs: set[Path]) -> list[Path]:
    """PDFs without a processable .docx sibling (still used for appendix splits)."""
    orphans: list[Path] = []
    for pdf in sorted(input_dir.glob("*.pdf")):
        if pdf.resolve() in paired_pdfs:
            continue
        orphans.append(pdf)
    return orphans


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input-dir",
        type=Path,
        action="append",
        default=None,
        help="Folder of paired .docx + .pdf files (repeatable; "
        "default: Inputs/DocumentImporter)",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=DEFAULT_OUT_DIR,
        help="Output root (default: Output/06_DocumentImporter)",
    )
    parser.add_argument(
        "--src",
        type=Path,
        default=None,
        help="Optional single .docx (uses sibling .pdf unless --pdf is set)",
    )
    parser.add_argument("--pdf", type=Path, default=None)
    parser.add_argument(
        "--appendices-only",
        action="store_true",
        help="Only extract multi-page appendix PDFs (skip DOCX/MD conversion)",
    )
    parser.add_argument(
        "--appendix-headers-only",
        action="store_true",
        help="In MD/DOCX, keep appendix headings only (skip appendix paras/tables); "
        "default includes full appendix content. Appendix PDF splits are unchanged.",
    )
    args = parser.parse_args(argv)

    if args.src is not None:
        src = args.src
        if not src.exists():
            print(f"Source not found: {src}", file=sys.stderr)
            return 1
        pdf = args.pdf or src.with_suffix(".pdf")
        if not pdf.exists():
            print(f"PDF not found: {pdf}", file=sys.stderr)
            return 1
        if args.appendices_only:
            extract_appendix_pdfs(pdf, args.out_dir, stem=src.stem)
        else:
            extract(
                src,
                args.out_dir,
                pdf,
                appendix_headers_only=args.appendix_headers_only,
            )
        return 0

    input_dirs = args.input_dir or list(DEFAULT_INPUT_DIRS)
    failures = 0
    any_work = False

    for input_dir in input_dirs:
        if not input_dir.is_dir():
            print(f"Input directory not found: {input_dir}", file=sys.stderr)
            failures += 1
            continue

        pairs = discover_pairs(input_dir)
        paired_pdfs = {pdf.resolve() for _, pdf in pairs}
        orphans = discover_orphan_pdfs(input_dir, paired_pdfs)

        if not pairs and not orphans:
            print(f"No PDF/DOCX work found in {input_dir}", file=sys.stderr)
            continue

        for docx, pdf in pairs:
            any_work = True
            print(f"\n=== {docx.stem} ===")
            try:
                if args.appendices_only:
                    extract_appendix_pdfs(pdf, args.out_dir, stem=docx.stem)
                else:
                    extract(
                        docx,
                        args.out_dir,
                        pdf,
                        appendix_headers_only=args.appendix_headers_only,
                    )
            except Exception as exc:  # noqa: BLE001
                failures += 1
                print(f"Error processing {docx.name}: {exc}", file=sys.stderr)

        for pdf in orphans:
            any_work = True
            print(f"\n=== {pdf.stem} (PDF appendices only) ===")
            try:
                extract_appendix_pdfs(pdf, args.out_dir, stem=pdf.stem)
            except Exception as exc:  # noqa: BLE001
                failures += 1
                print(f"Error extracting appendices from {pdf.name}: {exc}", file=sys.stderr)

    if not any_work:
        print("No documents processed", file=sys.stderr)
        return 1
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
